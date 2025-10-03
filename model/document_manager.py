"""Utilities for handling screenplay documents with python-docx."""
from __future__ import annotations

import os
import platform
import subprocess
from pathlib import Path
from typing import Optional

try:
    from docx import Document
except ImportError:  # pragma: no cover - python-docx might not be installed in CI
    Document = None  # type: ignore


class DocumentManager:
    """Create and open screenplay documents stored as .docx files."""

    @staticmethod
    def create_document(path: str | os.PathLike[str], title: str = "", summary: Optional[str] = None) -> None:
        """Create a new .docx document with optional title and summary."""
        if Document is None:
            raise RuntimeError("python-docx is required to create documents.")

        doc_path = Path(path)
        doc_path.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        if title:
            document.add_heading(title, level=1)
        if summary:
            document.add_paragraph(summary)
        document.save(doc_path)

    @staticmethod
    def ensure_document(path: str | os.PathLike[str], title: str = "", summary: Optional[str] = None) -> None:
        doc_path = Path(path)
        if not doc_path.exists():
            DocumentManager.create_document(doc_path, title=title, summary=summary)

    @staticmethod
    def open_doc(path: str | os.PathLike[str]) -> None:
        doc_path = Path(path)
        if not doc_path.exists():
            return

        system = platform.system()
        if system == "Windows":
            os.startfile(str(doc_path))  # type: ignore[attr-defined]
        elif system == "Darwin":
            subprocess.Popen(["open", str(doc_path)])
        else:
            subprocess.Popen(["xdg-open", str(doc_path)])
